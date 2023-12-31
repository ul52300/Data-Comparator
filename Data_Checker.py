import PySimpleGUI as sg
import pandas as pd
import os
import math
import datetime as dt
from docx import *
from copy import deepcopy
# Program: Data_Checker.py
# Version: 1.0.7
# Description: This program is used in order to display data onto a GUI inorder for the data to be compared to see if they match or not.
# Functions:
#   (1) To read Excel sheets and Microsoft Documents.
#   (2) Use parsed data in two tables that will be displayed on a GUI.

# Settings for the window.
window = sg.FlexForm('Data Checker', default_button_element_size = (5,2), auto_size_buttons=False, grab_anywhere=False, resizable=False)

# Column global variables.
COL_HEADINGS = ["Plot #", "RF Exposure Condition", "Mode", "Test Position", "Ch #.", "Freq. (MHz)", "RB Allocation", "RB Offset", "Max Area (W/kg)", "1-g Meas. (W/kg)", "10-g Meas. (W/kg)", "Peak SAR Location (x,y,z)", "Power Drift (dB)"]
COL_HEADINGS_LIQUID = ["Plot #", "SAR Lab", "Date & Time", "Frequency (MHz)", "Permittivity on Plot", "Conductivity on Plot", "Permittivity on PRN", "Conductivity on PRN", "Match?"]
COL_HEADINGS_COMPARATOR = ["Plot #", "RF Exposure Condition", "Mode", "Test Position",  "Match?", "If not, what is error?"]
COL_HEADINGS_EQUIPMENT_PROBE =  ["Plot #", "SAR Lab", "Probe SN", "Probe Cal Date", "Probe Cal Due Date"]
COL_HEADINGS_EQUIPMENT_DAE = ["Plot #", "SAR Lab", "DAE SN", "DAE Cal Date", "DAE Cal Due Date"]
COL_HEADINGS_PWR_DRIFT = ["Plot #", "Power Drift (dB)", "Within ±0.2 dB?"]
COL_WIDTHS = [len(COL_HEADINGS[0]), len(COL_HEADINGS[1])-4, len(COL_HEADINGS[2])+16, len(COL_HEADINGS[3]), len(COL_HEADINGS[4])+5, len(COL_HEADINGS[5]), len(COL_HEADINGS[6]), len(COL_HEADINGS[7]), len(COL_HEADINGS[8]), len(COL_HEADINGS[9]), len(COL_HEADINGS[10]), len(COL_HEADINGS[11]), len(COL_HEADINGS[12])]
COL_WIDTHS_COMPARATOR = [len(COL_HEADINGS_COMPARATOR[0])-5, len(COL_HEADINGS_COMPARATOR[1])-4, len(COL_HEADINGS_COMPARATOR[2])+16, len(COL_HEADINGS_COMPARATOR[3])-5, len(COL_HEADINGS_COMPARATOR[4]), len(COL_HEADINGS_COMPARATOR[5])]
COL_WIDTHS_EQUIPMENT = [len(COL_HEADINGS_EQUIPMENT_PROBE[0])-5, len(COL_HEADINGS_EQUIPMENT_PROBE[1]), len(COL_HEADINGS_EQUIPMENT_PROBE[2]), len(COL_HEADINGS_EQUIPMENT_PROBE[3]), len(COL_HEADINGS_EQUIPMENT_PROBE[4])]
COL_WIDTHS_PWR_DRIFT = [len(COL_HEADINGS_PWR_DRIFT[index]) for index in range(0, len(COL_HEADINGS_PWR_DRIFT))]

# Font global variables.
NORMAL_FONT = ("Times New Roman", 10)
NORMAL_OUTPUT_FONT = ("Times New Roman", 10, "bold")
TABLE_FONT = ("Times New Roman", 10, "bold")
BUTTON_FONT = ("Times New Roman", 10, "bold")
TABLE_HEADER_FONT = ("Times New Roman", 14, "bold")

# Text size:
DEFAULT_TEXT_SIZE = (20,1)
CHOOSE_TEXT_SIZE = (20,1)
ERROR_TEXT_SIZE = (20,1)
EXTREMITY_TEXT_SIZE = (20,1)
TABLE_HEADER_TEXT_SIZE = (7,2)

# Button size:
DEFAULT_BUTTON_SIZE = (10,1)
LOAD_BUTTON_SIZE = (10,1)
BROWSE_BUTTON_SIZE = (10,1)
CONFIRM_BUTTON_SIZE = (10,1)
EQUIPMENT_BUTTON_SIZE = (20,1)

# Input box size:
INPUT_EXCELDOCX_SIZE = (100,1)
TECHNOLOGY_COMBO_SIZE = (25,1)
EXTREMITY_COMBO_SIZE = (25,1)

# Number of rows for tables.
TABLE_NUM_ROWS = 9

# Flag for when hide button is pressed.
HIDE_COUNTER_MAIN = 0
HIDE_COUNTER_COMPARATOR = 0

# This is the main window.
def make_win1():
    # Theme
    sg.theme('NeutralBlue')
    
    # The layout of the window.
    layout = [ # The browser for choosing an excel file to be parsed.
            [sg.Text("Choose an excel file:",
                     size=CHOOSE_TEXT_SIZE,
                     font=NORMAL_FONT),
            sg.Input(key="-data_1-",
                     readonly=True,
                     size=INPUT_EXCELDOCX_SIZE),
            sg.FileBrowse("Browse",
                          file_types=(("Excel Files", "*.xlsx"),),
                          size=BROWSE_BUTTON_SIZE,
                          font=BUTTON_FONT,
                          tooltip="Choose the desired xlsx file that has the data."),
            sg.Text("Need Peak Location?:",
                     size=CHOOSE_TEXT_SIZE,
                     font=NORMAL_FONT),
            sg.Combo(values=["Yes", "No"],
                     key="-peak_sar-",
                     size=(5,1),
                     default_value="Yes")],
            
            # Once an excel is loaded, pick a sheet (technology) that will be parsed.
            [sg.Text("Pick a Technology:", 
                     size=CHOOSE_TEXT_SIZE, 
                     font=NORMAL_FONT),
            sg.Combo(values="", 
                     key="-tech_1-",
                     readonly=True,
                     size=TECHNOLOGY_COMBO_SIZE),
            sg.Button("Load",
                    size=LOAD_BUTTON_SIZE,
                    font=BUTTON_FONT,
                    tooltip="Refreshes technology list.\nHit this after loading a new excel file."),
            sg.Text("",
                    key="-Error_Technology-",
                    size=ERROR_TEXT_SIZE,
                    font=NORMAL_FONT)],
                        
            # The browser for choosing a docx file to be parsed.
            [sg.Text("Choose a docx file:",
                     size=CHOOSE_TEXT_SIZE,
                     font=NORMAL_FONT),
            sg.Input(key="-data_2-",
                     readonly=True,
                     size=INPUT_EXCELDOCX_SIZE),
            sg.FileBrowse("Browse",
                          file_types=(("Microsoft Document", "*.docx"),),
                          size=BROWSE_BUTTON_SIZE,
                          font=BUTTON_FONT,
                          tooltip="Choose the desired docx file that has the data.")],
            
            # Input on whether extremity is used or not. (NOTE: Currently this is hidden, because we don't delete rows in the excels normally, so the amount of rows in the excel that is being parsed should not change)
            [sg.Text("Extremity?:",
                     size=EXTREMITY_TEXT_SIZE,
                     font=NORMAL_FONT,
                     visible=False),
            sg.Combo(values=["Yes", "No"],
                     default_value="Yes", 
                     key="-confirm_extremity_1-", 
                     size=EXTREMITY_COMBO_SIZE,
                     disabled=True,
                     visible=False)], 
            
            [sg.HorizontalSeparator()],
            
            # Table for showing the data from the excel.
            [sg.VerticalSeparator(),
            sg.Text("Excel\nData:",
                    size=TABLE_HEADER_TEXT_SIZE,
                    font=TABLE_HEADER_FONT,
                    text_color="green"),
            sg.Table(values="", 
                     headings=COL_HEADINGS,  
                     key="-data_table_1-",
                     justification='center',
                     col_widths=COL_WIDTHS,
                     num_rows=TABLE_NUM_ROWS,
                     font=TABLE_FONT,
                     alternating_row_color="grey50",
                     row_colors=None,
                     expand_x=True,
                     expand_y=True,
                     auto_size_columns=False,
                     vertical_scroll_only=False,
                     enable_events=True)],
            
            [sg.HorizontalSeparator()],
            
            # Table for showing the data from the plots.
            [sg.VerticalSeparator(),
            sg.Text("Plot\nData:",
                    size=TABLE_HEADER_TEXT_SIZE,
                    font=TABLE_HEADER_FONT,
                    text_color="blue"),
            sg.Table(values="", 
                     headings=COL_HEADINGS,  
                     key="-data_table_2-",
                     justification='center',
                     col_widths=COL_WIDTHS,
                     num_rows=TABLE_NUM_ROWS,
                     font=TABLE_FONT,
                     alternating_row_color="grey50",
                     expand_x=True,
                     expand_y=True,
                     auto_size_columns=False,
                     vertical_scroll_only=False,
                     enable_events=True)],
            
            [sg.HorizontalSeparator()],
            
            # The buttons for: loading an excel, loading a docx, and compare results window.
            [sg.Button("Load Excel",
                       size=CONFIRM_BUTTON_SIZE,
                       font=BUTTON_FONT,
                       button_color="white",
                       auto_size_button=False,
                       tooltip="Press this to load the data from the Excel sheet."),
            sg.Button("Load Docx",
                      size=CONFIRM_BUTTON_SIZE,
                      font=BUTTON_FONT,
                      button_color="white",
                      tooltip="Press this to load the data from the plot."),
            sg.Button("Compare",
                      key="-compare-",
                      size=CONFIRM_BUTTON_SIZE,
                      font=BUTTON_FONT,
                      button_color="white",
                      tooltip="Press this to open a window to compare the 1-g and 10-g data from both tables."),
            sg.Button("Equipment",
                     key="-equipment-",
                     size=CONFIRM_BUTTON_SIZE,
                     font=BUTTON_FONT,
                     button_color="white",
                     tooltip="Press this to open the equipment menu."),
            sg.Button("Liquid Checker",
                      key="-liquid_check-",
                      size=(12,1),
                      font=BUTTON_FONT,
                      button_color="white",
                      tooltip="Press this to open the Liquid Checker"),
            sg.Text("",
                    key="-FuckedUp-",
                    size=(20,1),
                    font=BUTTON_FONT)],
            
            [sg.Button("Hide/Unhide",
                       key="-hide_unhide-",
                       size=(12,1),
                       font=BUTTON_FONT,
                       button_color="white",
                       auto_size_button=False,
                       tooltip="Press to hide/unhide specified columns."),
             sg.Combo(values=["Show all", "Only 1g/10g values", "No 1g/10g values"],
                      default_value="",
                      key="-hide_unhide_list-",
                      size=(20,1),
                      readonly=True)]
    ]
    return(sg.Window("Data Checker", layout, size=(1200,530), resizable=True, return_keyboard_events=True, location=(0,0), finalize=True))   # Display the window.

# This is the liquid check window.
def make_win2():
    sg.theme('NeutralBlue')
    layout = [
            [sg.Text("Notes about this window:",
                     font=NORMAL_OUTPUT_FONT)],
            [sg.Text("- Use the top portion of this window if you just want to know the permittivity and conductivtiy of a single target frequency.",
                     font=NORMAL_FONT)],
            [sg.Text("- Use the bottom portion if you want to check the permittivity and conductivity of a WHOLE plot.",
                     font=NORMAL_FONT)],
            [sg.HorizontalSeparator()],
            [sg.Text("Choose a .prn file from SAR drive:", 
                    size=(26,1), 
                    font=NORMAL_FONT), 
            sg.Input(
                key="-file_1-", 
                size=(10,1),
                font=BUTTON_FONT,
                readonly=True), 
            sg.FileBrowse(size=(10,1),
                          initial_folder="\\\\FREshares\\SAR\\5. Liquid Check (result and prn)")],
            [sg.Text("Input a target (MHz):", 
                    size=(26,1), 
                    font=NORMAL_FONT), 
            sg.InputText(key="-target_1-", 
                        size=(10,1))],
            [sg.Txt('')],
            [sg.Text("Results", 
                     font=("Times New Roman", 12, "bold", "underline"))],
            [sg.Text("Target Frequency (MHz):", 
                     size=(20,1), 
                     font=('Times New Roman', 10)), 
            sg.Text('',
                    key='input_1',
                    size=(10, 1), 
                    font=NORMAL_OUTPUT_FONT)],
            [sg.Text("Permittivity:", 
                     size=(20,1), 
                     font=('Times New Roman', 10)), 
            sg.Text('',
                    key='input_2',
                    size=(10, 1), 
                    font=NORMAL_OUTPUT_FONT)],
            [sg.Text("Conductivity:", 
                     size=(20,1), 
                     font=('Times New Roman', 10)), 
            sg.Text('',
                    key='input_3',
                    size=(10, 1), 
                    font=NORMAL_OUTPUT_FONT)],
            [sg.Txt('')],
            [sg.Button("Calculate", 
                       size=(10,1),
                       font=BUTTON_FONT)],
            [sg.HorizontalSeparator()],
            [sg.Text("Plot's Liquid Parameters",
                     font=("Times New Roman", 12, "bold", "underline"))],
            [sg.Table(values="", 
                      headings=COL_HEADINGS_LIQUID,  
                      key="-liquid_table_1-",
                      justification='center',
                      col_widths=[],
                      num_rows=TABLE_NUM_ROWS,
                      font=TABLE_FONT,
                      alternating_row_color="grey50",
                      expand_x=True,
                      expand_y=True,
                      auto_size_columns=False,
                      vertical_scroll_only=False,
                      enable_events=True)],
            [sg.Button("Load Parameters",
                       size=(15,1),
                       font=BUTTON_FONT),
            sg.Push(), 
            sg.Button("Quit", 
                      size=(10,1),
                      font=BUTTON_FONT)]
    ]
    return(sg.Window("Liquid Checker", layout, resizable=True, location=(0,0), finalize=True))  # Display the window.

# This is the comparator window.
def make_win3():
    sg.theme('NeutralBlue') # Theme
    layout = [
             [sg.Table(values="",
                    headings=COL_HEADINGS_COMPARATOR,
                    key="-data_table_3-",
                    justification='center',
                    col_widths=COL_WIDTHS_COMPARATOR,
                    num_rows=18,
                    font=TABLE_FONT,
                    expand_x=True,
                    expand_y=True,
                    auto_size_columns=True,
                    vertical_scroll_only=False,
                    enable_events=True)],
            [sg.Button("Prepare to Be Sad",
                    size=(30,1),
                    font=BUTTON_FONT,
                    tooltip="Press to compare the data.")]]
    return(sg.Window("Sadness?", layout, size=(1200, 350), resizable=True, location=(0,0), finalize=True))

# This is the equipment window.
def make_win4():
    sg.theme('NeutralBlue') # Theme
    layout = [
             [sg.Table(values="",
                       headings=COL_HEADINGS_EQUIPMENT_PROBE,
                       key="-equipment_table_1-",
                       justification='center',
                       col_widths=COL_WIDTHS_EQUIPMENT,
                       num_rows=9,
                       font=TABLE_FONT,
                       alternating_row_color="grey50",
                       expand_x=True,
                       expand_y=True,
                       auto_size_columns=False,
                       enable_events=True)],
             [sg.Table(values="",
                       headings=COL_HEADINGS_EQUIPMENT_DAE,
                       key="-equipment_table_2-",
                       justification='center',
                       col_widths=COL_WIDTHS_EQUIPMENT,
                       num_rows=9,
                       font=TABLE_FONT,
                       alternating_row_color="grey50",
                       expand_x=True,
                       expand_y=True,
                       auto_size_columns=False,
                       enable_events=True)],
             [sg.Button("Load Equipment",
                        key="-load_equipment-",
                        size=EQUIPMENT_BUTTON_SIZE,
                        font=BUTTON_FONT,
                        tooltip="Press to load in equipment.")]]
    return(sg.Window("Equipment Menu", layout, size=(670,360), resizable=True, location=(0,0), finalize=True))    

# This function's purpose is to get a PRN file based on the SAR Lab and date of the current plot.
def find_prn(target, sar_lab, date_prn, months):
    # Defining each building's path.
    building_one_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 1 SAR Labs")
    building_twofront_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 2 Front SAR Labs")
    building_tworear_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 2 Rear SAR Labs")

    # Logic for determining where the SAR Lab is located.
    if sar_lab in building_one_dir:
        hsl_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 1 SAR Labs\\{}\\prn data for (for importing to DASY medium parameters)\\Head".format(sar_lab))
        building = "Building 1 SAR Labs"
    elif sar_lab in building_twofront_dir:
        hsl_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 2 Front SAR Labs\\{}\\prn data for (for importing to DASY medium parameters)\\Head".format(sar_lab))
        building = "Building 2 Front SAR Labs"
    elif sar_lab in building_tworear_dir:
        hsl_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\Building 2 Rear SAR Labs\\{}\\prn data for (for importing to DASY medium parameters)\\Head".format(sar_lab))
        building = "Building 2 Rear SAR Labs"
        
    # Logic for determining the folder you get for frequency.
    closest_freq = []
    for hsl in hsl_dir:
        if any(liquid in hsl for liquid in ["3-250", "4-250", "30-250"]):
            continue
        elif len(closest_freq) == 0:
            closest_freq.append("6")
            closest_freq.append(abs(float(target)-6))
        elif hsl[3:] == "5000" and (target >= 5000 and target < 6000):
            closest_freq[0] = "5000"
            break
        else:
            freq_distance = abs(float(target)-float(hsl[3:]))
            if freq_distance <= closest_freq[1]:
                closest_freq[0] = hsl[3:]
                closest_freq[1] = freq_distance
    
    # Logic to get the correct PRN file. (NOTE: 4 days is the maximum days allowed for a valid liquid check).
    hsl_freq = "HSL" + closest_freq[0]    
    prn_dir = os.listdir("\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\{}\\{}\\prn data for (for importing to DASY medium parameters)\\Head\\{}".format(building, sar_lab, hsl_freq))
    found = 0
    count_days = 0
    
    while found != 1 or count_days < 5:
        index_prnfolder = 0
        for prn in prn_dir:
            if date_prn in prn:
                prn_found = prn_dir[index_prnfolder]
                found = 1
                break
            index_prnfolder += 1
        
        if found == 1 or count_days >= 5:
            break
        
        date_prn = dt.datetime.strptime(date_prn, "%Y-%b-%d")  
        date_prn = date_prn - dt.timedelta(days=1)
        date_prn = date_prn.strftime("%Y-%b-%d")
        count_days += 1
    
    prn_file = "\\\\FREshares\\SAR\\5. Liquid Check (result and prn)\\{}\\{}\\prn data for (for importing to DASY medium parameters)\\Head\\{}\\{}".format(building, sar_lab, hsl_freq, prn_found)    
    return(prn_file)

# This function's purpose is to calculate the permittivity and conductivity of a particular target frequency.
def permcondCalcSolo(target, lines):
    low_freq = target - (target % 5) - 5    # Getting the 'low' freq from the target.
    high_freq = target - (target % 5) + 5 # Getting the 'high' freq from the target.
    
    # Logic to get the relative permitivity and relative conductivity from .prn file.
    for line in lines:
        if line.find(str(int(low_freq * pow(10,6)))) != -1:
            rperm_cond_low = list((line.rstrip().split()[1],line.rstrip().split()[2]))
        if line.find(str(int(high_freq * pow(10,6)))) != -1:
            rperm_cond_high = list((line.rstrip().split()[1],line.rstrip().split()[2]))
            break
    
    if target < 20 or target > 7200: # The .prn file does not go below 20 MHz and above 7200 MHz.
        window['input_1'].update("N/A")
        window['input_2'].update("N/A")
        window['input_3'].update("N/A")    
    else:
        # Logic for linear interpolation of the relative permitivity and conductivity for the target frequency.
        rpermitivity = ((high_freq - target)/(high_freq - low_freq))* float(rperm_cond_low[0]) + ((low_freq - target)/(low_freq - high_freq)) * float(rperm_cond_high[0])
        rconductivity = ((high_freq - target)/(high_freq - low_freq))* float(rperm_cond_low[1]) + ((low_freq - target)/(low_freq - high_freq)) * float(rperm_cond_high[1])
        conductivity = rconductivity * (2 * math.pi) * (target * pow(10,6)) * (8.854 * pow(10,-12))
        window['input_1'].update(target)
        window['input_2'].update(round(rpermitivity,1))
        window['input_3'].update(round(conductivity,3))

# This function's purpose is to calculate the permittivity and conductivity so that it will be compared to the plot's data.
def permcondCalcTable(target, lines):
    low_freq = target - (target % 5) - 5    # Getting the 'low' freq from the target.
    high_freq = target - (target % 5) + 5 # Getting the 'high' freq from the target.
        
    # Logic to get the relative permitivity and relative conductivity from .prn file.
    for line in lines:
        if line.find(str(int(low_freq * pow(10,6)))) != -1:
            rperm_cond_low = list((line.rstrip().split()[1],line.rstrip().split()[2]))
        if line.find(str(int(high_freq * pow(10,6)))) != -1:
            rperm_cond_high = list((line.rstrip().split()[1],line.rstrip().split()[2]))
            break
    
    if target < 20 or target > 7200: # The .prn file does not go below 20 MHz and above 6000 MHz.
        rpermitivity = "N/A"
        conductivity = "N/A"  
    else:
        # Logic for linear interpolation of the relative permitivity and conductivity for the target frequency.
        rpermitivity = ((high_freq - target)/(high_freq - low_freq))* float(rperm_cond_low[0]) + ((low_freq - target)/(low_freq - high_freq)) * float(rperm_cond_high[0])
        rconductivity = ((high_freq - target)/(high_freq - low_freq))* float(rperm_cond_low[1]) + ((low_freq - target)/(low_freq - high_freq)) * float(rperm_cond_high[1])
        conductivity = rconductivity * (2 * math.pi) * (target * pow(10,6)) * (8.854 * pow(10,-12))
    return(str(rpermitivity), str(conductivity))

# This function's purpose is to compare the data from both the Excel and plots.
def append_data(match_list, excel, plot):
    compare_list = []
    
    for data_1, data_2, index in zip(excel, plot, range(0, len(data_excel))):
        rf_exposure_cond_excel, rf_exposure_cond_plot = data_1[1], data_2[1]   # Assign RF Exposure Conditions from both excel and plot.
        mode_excel, mode_plot =                         data_1[2], data_2[2]   # Assign test mode from both excel and plot.
        test_position_excel, test_position_plot =       data_1[3], data_2[3]   # Assign Test Position from both excel and plot.
        channel_num_excel, channel_num_plot =           data_1[4], data_2[4]   # Assign Channel # from both excel and plot.
        frequency_excel, frequency_plot =               data_1[5], data_2[5]   # Assign Frequency (MHz) from both excel and plot.
        rb_allocation_excel, rb_allocation_plot =       data_1[6], data_2[6]   # Assign RB Allocation from both excel and plot.
        rb_offset_excel, rb_offset_plot =               data_1[7], data_2[7]   # Assign RB Offset from both excel and plot.
        max_area_scan_1g_excel, max_area_scan_1g_plot = data_1[8], data_2[8]   # Assign Max Area Scan 1g (W/kg) from both excel and plot.
        one_g_meas_excel, one_g_meas_plot =             data_1[9], data_2[9]   # Assign Measured 1g SAR (W/kg) from both excel and plot.
        ten_g_meas_excel, ten_g_meas_plot =             data_1[10], data_2[10] # Assign Measured 10g SAR (W/kg) from both excel and plot.
        
        # Order the different parameters into a list for comparison.
        excel_data = [
                        rf_exposure_cond_excel,
                        mode_excel,
                        test_position_excel,
                        channel_num_excel,
                        frequency_excel,
                        rb_allocation_excel,
                        rb_offset_excel,
                        max_area_scan_1g_excel,
                        one_g_meas_excel,
                        ten_g_meas_excel]
        # Order the different parameters into a list for comparison.
        plot_data = [
                        rf_exposure_cond_plot,
                        mode_plot,
                        test_position_plot,
                        channel_num_plot,
                        frequency_plot,
                        rb_allocation_plot,
                        rb_offset_plot,
                        max_area_scan_1g_plot,
                        one_g_meas_plot,
                        ten_g_meas_plot]

        # Logic for the comparator.
        match_list.append([index + 1]) # Insert plot number.
        for excel_data_index, plot_data_index in zip(excel_data, plot_data):
            # Insert a 'Y' if the data matches or a 'N' if it doesn't. No comparison needed for first 4 columns, so add it to table.
            if excel_data_index in excel_data[0:3] and plot_data_index in plot_data[0:3] and excel_data_index.lower() == plot_data_index.lower():
                match_list[index] = match_list[index] + [excel_data_index]
            elif excel_data_index == plot_data_index:
                match_list[index] = match_list[index] + ["Y"]
            else:
                match_list[index] = match_list[index] + ["N"]

    # Logic to tell what is not lining up in for data comparisons.
    for row, index in zip(match_list, range(0, len(match_list))):
        compare_list.append(row[:4])
        if "N" in row[4:]:
            compare_list[index].append("N")
            
            indices = []
            for iter in range(0, len(row[4:])):
                if row[4:][iter] == "N":
                    indices.append(iter+4)
                    
            for error_col in indices:
                if len(compare_list[index]) < 6:
                    compare_list[index].append(COL_HEADINGS[error_col])
                else:
                    compare_list[index][5] = compare_list[index][5] + ", " + COL_HEADINGS[error_col]
        else:
            compare_list[index].append("Y")
            compare_list[index].append("No Error")
        
    return(match_list, compare_list)
    
xl = path = tech = ""
data_excel = []
data_plot = []
skip_rows, num_rows = 0, 0

rb_positions = {
    "1400000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "3",
                "High": "5"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "1",
                "High": "3"
            }
        }
    },
    "3000000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "8",
                "High": "14"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "4",
                "High": "7"
            }
        }
    },
    "5000000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "12",
                "High": "24"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "7",
                "High": "13"
            }
        },
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "12",
                    "High": "23"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "6",
                    "High": "7"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "5",
                    "High": "9"
                },
                "50% RB": {
                    "Low": "2",
                    "Mid": "3",
                    "High": "4"
                }                
            }  
        }
    },
    "10000000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "25",
                "High": "49"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "12",
                "High": "25"
            }
        },
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "26",
                    "High": "50"
                },
                "50% RB": {
                    "Low": "12",
                    "Mid": "14",
                    "High": "15"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "12",
                    "High": "22"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "6",
                    "High": "--"
                }                
            }          
        }
    },
    "15000000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "37",
                "High": "74"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "20",
                "High": "39"
            }
        },
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "39",
                    "High": "77"
                },
                "50% RB": {
                    "Low": "18",
                    "Mid": "22",
                    "High": "25"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "19",
                    "High": "36"
                },
                "50% RB": {
                    "Low": "9",
                    "Mid": "10",
                    "High": "11"
                }                
            }        
        }
    },
    "20000000": {
        "LTE": {
            "1 RB": {
                "Low": "0",
                "Mid": "49",
                "High": "99"
            },
            "50% RB": {
                "Low": "0",
                "Mid": "24",
                "High": "50"
            }
        },
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "53",
                    "High": "104"
                },
                "50% RB": {
                    "Low": "25",
                    "Mid": "28",
                    "High": "31"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "25",
                    "High": "49"
                },
                "50% RB": {
                    "Low": "12",
                    "Mid": "13",
                    "High": "14"
                }                
            }        
        }
    },
    "25000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "66",
                    "High": "131"
                },
                "50% RB": {
                    "Low": "32",
                    "Mid": "35",
                    "High": "37"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "32",
                    "High": "63"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "16",
                    "High": "17"
                }                
            }        
        }
    },
    "30000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "80",
                    "High": "158"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "40",
                    "High": "--"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "39",
                    "High": "76"
                },
                "50% RB": {
                    "Low": "18",
                    "Mid": "21",
                    "High": "24"
                }                
            }        
        }
    },
    "35000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "94",
                    "High": "186"
                },
                "50% RB": {
                    "Low": "45",
                    "Mid": "49",
                    "High": "53"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "46",
                    "High": "90"
                },
                "50% RB": {
                    "Low": "22",
                    "Mid": "24",
                    "High": "25"
                }                
            }        
        }
    },
    "40000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "108",
                    "High": "214"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "54",
                    "High": "--"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "53",
                    "High": "104"
                },
                "50% RB": {
                    "Low": "25",
                    "Mid": "28",
                    "High": "31"
                }                
            }     
        }
    },
    "45000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "121",
                    "High": "240"
                },
                "50% RB": {
                    "Low": "60",
                    "Mid": "61",
                    "High": "62"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "59",
                    "High": "117"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "30",
                    "High": "--"
                }                
            }
        }
    },
    "50000000": {
        "NR": {
            "15000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "135",
                    "High": "268"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "67",
                    "High": "68"
                }
            },
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "66",
                    "High": "131"
                },
                "50% RB": {
                    "Low": "32",
                    "Mid": "35",
                    "High": "37"
                }                
            }
        }
    },
    "60000000": {
        "NR": {
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "81",
                    "High": "160"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "40",
                    "High": "41"
                }                
            }
        }
    },
    "70000000": {
        "NR": {
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "94",
                    "High": "187"
                },
                "50% RB": {
                    "Low": "45",
                    "Mid": "50",
                    "High": "54"
                }                
            }
        }
    },
    "80000000": {
        "NR": {
            "30": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "108",
                    "High": "215"
                },
                "50% RB": {
                    "Low": "--",
                    "Mid": "54",
                    "High": "55"
                }                
            }
        }
    },
    "90000000": {
        "NR": {
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "122",
                    "High": "243"
                },
                "50% RB": {
                    "Low": "60",
                    "Mid": "63",
                    "High": "65"
                }                
            }
        }
    },
    "100000000": {
        "NR": {
            "30000": {
                "1 RB": {
                    "Low": "1",
                    "Mid": "136",
                    "High": "271"
                },
                "50% RB": {
                    "Low": "67",
                    "Mid": "69",
                    "High": "71"
                }                
            }
        }
    }
}

nr_nrbs = {
    "15000": {
        "5000000": "25",
        "10000000": "50",
        "15000000": "75",
        "20000000": "100",
        "25000000": "128",
        "30000000": "160",
        "35000000": "180",
        "40000000": "216",
        "45000000": "242",
        "50000000": "270"
    },
    "30000": {
        "5000000": "11",
        "10000000": "24",
        "15000000": "38",
        "20000000": "51",
        "25000000": "64",
        "30000000": "75",
        "35000000": "90",
        "40000000": "100",
        "45000000": "119",
        "50000000": "128",
        "60000000": "162",
        "70000000": "180",
        "80000000": "216",
        "90000000": "243",
        "100000000": "270"        
    }
}
window_main, window_liquid, window_compare, window_equipment = make_win1(), None, None, None
while True:
    window, event, values = sg.read_all_windows()
    
    # Break out of loop which closes the window(s).
    if event == sg.WIN_CLOSED or event == 'Quit':
        window.close()
        if window == window_liquid:
            window_liquid = None
        elif window == window_compare:
            window_compare = None
        elif window == window_equipment:
            window_equipment = None
        elif window == window_main:
            break
    
    if window == window_main and event != sg.WIN_CLOSED:
        window["-FuckedUp-"].update("")
        path_excel = values["-data_1-"]                         # This is the path used to get the excel.
        path_docx = values["-data_2-"]                          # This is the path used to get the microsoft document.
        tech = values["-tech_1-"]                               # This is the technology/band that will be selected.
        extremity_confirm = values["-confirm_extremity_1-"]     # This is a confirmation for whether you need extremity or not.
        # excluded_exposure =  values["-excluded_exposure_1-"]  # This is the selection of body exposure condition. (Only Head and Body).
        # excluded_number = values["-excluded_number_1-"]       # This is the number of excluded positions due to distance being to far from the antenna.
       
    print("Event:", event)
    print("Technology:", tech)
    
    # When the "Load" button is pressed:
    # (1) Sheet names will be taken from the xlsx file.
    # (2) Remove unnecessary sheets from the sheet names list.
    # (3) Update the combo element that holds the sheet names.
    if event == "Load" and path_excel != "":
        
        # Clear window.
        window["-Error_Technology-"].update("")
        window["-FuckedUp-"].update("")
        
        try:
            # Put sheet names from Excel in list.
            xl = pd.ExcelFile(path_excel)
            xl_sheets = xl.sheet_names
        except:
            window["-Error_Technology-"].update("Please use an Excel file!")
            window["-FuckedUp-"].update("Dx")
        for sheet_name in ['How to Use this Workbook', 'Inter-Band CA Exclusion', 'Settings', 'Sum of SAR',
                        'Repeated', 'Data', 'Master List (WWAN)', 'List Variables', 'ISED Extremity']:
            if sheet_name in xl_sheets:
                xl_sheets.remove(sheet_name)
        window["-tech_1-"].update(values = xl_sheets)
        xl.close()
    elif event == "Load Excel" and tech != "":
        
        ref_df = pd.read_excel(path_excel, sheet_name=tech, index_col=None, na_values=['N/A'])
        ref_data = ref_df.values.tolist()
        
        # Get the parameters on where the parse the data in the excel.
        count, skip_rows, num_rows = 0, 0, 0
        for ref_row in ref_data:
            if ref_row[0] == "System Check Date" and skip_rows == 0:
                skip_rows = count + 1
                count = 0
            elif ref_row[0] == "Repeated" and skip_rows != 0:
                num_rows = count - 3
            count += 1

        del ref_df
        del ref_data
         
        '''
            These if-else statements are for assigning the columns that will be parsed in the Excel documents.
            
            Notes about this section:
            1) "tech": Is the Excel sheet that contains the name of the Band/Technology that was tested.
            2) "cols": Contains the column letters in the Excel document that have the data that will be parsed.
                For GSM, PCS, WCDMA:
                    M = Test Position
                    U = Channel #
                    V = Frequency
                    Y = 1-g Measured SAR
                    AA = 10-g Measured SAR
                For Wi-Fi:
                    N = Test Position
                    O = Channel #
                    P = Frequency
                    AA = Max Area Scan SAR
                    AD = 1-g Measured SAR
                    AF = 10-g Measured SAR
                For LTE:
                    M = Test Position
                    W = Channel #
                    X = Frequency
                    Y = RB Allocation
                    Z = RB Offset
                    AC = 1-g Measured SAR
                    AE = 10-g Measured SAR
                For Bluetooth:
                    N = Test Position
                    O = Channel #
                    P = Frequency
                    AC = 1-g Measured SAR
                    AE = 10-g Measured SAR
        '''
        if any(technology in tech for technology in ["GSM", "PCS", "W-CDMA"]):
            cols = "I:J, M, U:V, Y, AA, AC:AD"
            group = "GSM" if any(technology in tech for technology in ["GSM", "PCS"]) else "WCDMA"
        elif any(technology in tech.lower() for technology in ["wlan", "wi-fi", "wifi", "u-nii",]):
            cols = "I:J, N, O:P, AA, AD, AF, AH:AI"
            group = "WLAN"
        elif "LTE" in tech or "FR1" in tech:
            cols = "I:J, M, W:Z, AC, AE, AG:AH"
            group = "LTE" if "LTE" in tech else "FR1"
        elif "Bluetooth" in tech:
            cols = "I:J, N, O:P, AC, AE, AG:AH"
            group = "Bluetooth"
        
        df = pd.read_excel(path_excel, sheet_name=tech, index_col=None, na_values=['N/A'], usecols="{}".format(cols), skiprows=skip_rows, nrows=num_rows) # Create a dataframe from the excel on the selected rows and columns.
        df = df.fillna("N/A")       # Replace nan with N/A.
        data = df.values.tolist()   # Insert all the values of the dataframe into a list.        

        '''
            Notes about this section:
            This section will be what assigns the positions that each data column will be in. This makes it easier to tell where to put the data on the tables.
            
            # NOTE: Not all of these columns will have data in it. The columns that don't have data will instead have "N/A". (There will also be a "Hide/Unhide" button to hide whatever isn't needed).
            Order of the columns: [0 = RF Exposure Condition,
                                   1 = Mode,
                                   2 = Test Position
                                   3 = Frequency, 
                                   4 = Channel, 
                                   5 = RB Allocation (LTE, FR1) or Max Area Scan SAR (Wi-Fi), 
                                   6 = RB Offset (LTE, FR1) or 1-g Measured SAR (Wi-Fi) or 10-g Measured SAR (GSM, WCDMA),
                                   7 = 1-g Max Area Scan SAR (Wifi)
                                   8 = 1-g Measured SAR (LTE, FR1) or 10-g Measured SAR (Wi-Fi),
                                   9 = 10-g Measured SAR (LTE)]
            
            1) "merge_rf_exposure condition": This indicates the "first" column of the table, which is the rf exposure condition. It is called "merge" because, unless the rf exposure condition has only ONE channel, 
                                              the cell containing the rf exposure condition will be a merged cell (multiple cell). It is assigned to 0, since it will be the first column on the table (Not counting the Plot #).
            2) "merge_mode": This is the same as (1), except the mode is the information being parsed.
            3) "merge_test_position": This is the same as (1) and (2), except the test position being parsed.
            4) "max_area_scan": This holds the position of the max area scan SAR.
            5) "frequency_num" & "channel_num": This holds the frequency and channel position, respectively.
            6) "meas_1g" & "meas_10g": This holds the 1-g and 10-g measured SAR position, respectively.
            7) "freq_ch_nrb_orb": This holds the positions of the columns of a row. Goes from 1 to 6.
            8) "merge_variables": 
            9) "meas_values": This holds the positions of the measured SAR values and max area scan SAR.
        '''
        len_data = len(data)
        for position in range(0, len_data):
            # Initialize the flags for placing the data into a list.
            merge_rf_exposure_condition = 0     # Assign the flag for merged cells for the rf exposure position to 0.
            merge_mode = 1                      # Assign the flag for merged cells for the mode to 1.
            merge_test_position = 2             # Assign the flag for merged cells for the test position to 2.
            frequency_num = 3                   # Assign the flag for frequency to 3.
            channel_num = 4                     # Assign the flag for channel to 4.
            num_rb = 5                          # Assign the flag for RB allocation to 5. (NOTE: Only applies to LTE and FR1, otherwise these are ignored). 
            offset_rb = 6                       # Assign the flag for RB Offset to 6. (NOTE: Only applies to LTE and FR1, otherwise these are ignored).  
            max_area_scan = 5                   # Assign the flag for max area scan to 5. (NOTE: Only applies to Wi-Fi, otherwise this is ignored).
            
            # Assigning the positions for measured 1-g and measured 10-g. These are dependent on the technology.
            if ("WLAN" or "U-NII") in group:
                meas_1g = 6
                meas_10g = 7
                peak_sar = 8
                power_drift = 9
            elif "LTE" in group or "FR1" in group:
                meas_1g = 7
                meas_10g = 8
                peak_sar = 9
                power_drift = 10
            else:
                meas_1g = 5
                meas_10g = 6
                peak_sar = 7
                power_drift = 8
            
            freq_ch_nrb_orb = [frequency_num, channel_num, num_rb, offset_rb]
            merge_variables = [merge_rf_exposure_condition, merge_mode, merge_test_position]                
            meas_values = [max_area_scan, meas_1g, meas_10g]
            
            '''
                Notes about this section:
                This section is for placing the positions of the data that will be parsed into a list. By removing what is not needed for the columns for certain techonologies, 
                this makes the list that contains the positions more clean and organized.
            '''
            if any(technology in tech for technology in ["LTE", "FR1"]):
                meas_values.remove(max_area_scan)
            elif any(technology in tech for technology in ["GSM", "PCS", "W-CDMA", "Bluetooth"]):
                for not_needed in [num_rb, offset_rb]:
                    freq_ch_nrb_orb.remove(not_needed)
                meas_values.remove(max_area_scan)
            elif any(technology in tech for technology in ["WLAN", "WiFi", "Wi-Fi", "Wi-fi", "U-NII"]): # NOTE: You'll notice that Wifi has the same code as GSM and WCDMA. This is because, if I wanted to add duty cycle, I would need an additional slot for that.
                for not_needed in [num_rb, offset_rb]:
                    freq_ch_nrb_orb.remove(not_needed)
            '''
                Notes about this section:
                This section is for formatting the data from the dataframe list from the Excel document into whatever is needed.
                It works by iterating through the list of column positions, and based on the current column position will do certain things explained below.
                
                1) First 'if' statement:
                    Description: Excel has the ability to merge cells. When a cell is merged with another or more cells, this "merges" all selected cells. This also makes it so that the top-left most cell of the merged cell
                                 contains the actual data of the cell. The rest of the merged cells of what would be the individual cells are counted as blank or 'nan' (Not A Number). The problem to solve is that if I take the
                                 merged cell that actually takes the value, the next cell of the merged cell will be a blank. To solve this, as long as the next cell does not contain a different value from the 
                                 current value, it will copy the cell of the current value. However, once we find a different value, this will indicate that we are in fact on another value.
                2) Second 'if' statement:
                    Description: Insert the data for the frequency, channel, RB Allocation (if applicable), and RB Offset (if applicable).
                3) Third 'if' statement:
                    Description: Insert the 1-g max area scan, 1-g and 10-g measured SAR data. (NOTE: Specifically, these number have to be rounded to THREE decimal places, and show only THREE decimal places).
            '''            
            for sublist_number in range(0, len(data[position])):
                # Because of merged cells in the xlsx, the cells that are not at the top
                # of the merged cell are considered "NaN" ("N/A" since these were filled). This solves that by replacing
                # the index with a "NaN" ("N/A") with the previous index, which is the test position.
                if sublist_number in merge_variables and data[position][sublist_number] == "N/A":
                    data[position][sublist_number] = data[position-1][sublist_number]
                # Convert the "Ch. #" to an integer and add ""
                elif sublist_number in freq_ch_nrb_orb:
                    data[position][sublist_number] = "{}".format(data[position][sublist_number]) if (data[position][sublist_number] == "N/A" or isinstance(data[position][sublist_number], str)) else "{}".format(float(data[position][sublist_number]) if sublist_number == 4 else int(data[position][sublist_number]))
                # Round the "1-g Meas. (W/kg)" and "10-g Meas. (W/kg)" columns. (NOTE: Also rounds "Max Area (W/kg) column as well, if Wi-Fi")
                elif sublist_number in meas_values and data[position][sublist_number] != "N/A":
                    data[position][sublist_number] = "{:.3f}".format(round(data[position][sublist_number], 3))
                elif sublist_number == power_drift and data[position][sublist_number] != "N/A":
                    data[position][sublist_number] = "{:.2f}".format(round(data[position][sublist_number], 2))

            # "N/A"'s will be inserted into certain column positions if certain technologies are being parsed.
            if any(technology in group for technology in ["WLAN", "WCDMA", "GSM", "Bluetooth"]):
                data[position].insert(5, "N/A")
                data[position].insert(6, "N/A")
                if any(technology in group for technology in ["WCDMA", "GSM", "Bluetooth"]):
                    data[position].insert(7, "N/A")
            elif "LTE" in group or "FR1" in group:
                data[position].insert(7, "N/A")

        '''
            Notes about this section:
            This section detects if the current sublist has no number in the 1-g Meas or 10-g Meas (Or 'nan') and will remove a row if there is no max area scan, 1-g, and 10-g.
        '''
        len_data = len(data)
        plot_number_tracker = 0 # Used to update the index after a pop has happened.
        print(data)
        for index in range(0, len_data):
            index += plot_number_tracker
            if (data[index][len(data[index])-4] == "N/A" or data[index][len(data[index])-3] == "N/A") and (data[index][len(data[index])-5] == "N/A"
                                                                                                           or data[index][len(data[index])-5] == "nan" 
                                                                                                           or pd.isna(data[index][len(data[index])-5])):  
                data.pop(index) # Remove current index that contains 'nan' on 1-g Meas or 10-g Meas.
                len_data -= 1
                plot_number_tracker -= 1
            else:
                data[index].insert(0, index+1)
        
        df_excel = data       
        data_excel = df_excel # Used for comparison purposes.

        '''
            Notes about this section:
            The purpose of this section is to hide the columns on the table that are unnessesary when checking certain technologies. I.E. RB Allocation, RB Offset, and Max Area Scan.
        '''
        try:
            displaycolumns = deepcopy(COL_HEADINGS) # Creating a deepcopy as to not override the original.
            if any(technology in group for technology in ["GSM", "WCDMA", "WLAN", "Bluetooth"]):
                displaycolumns.remove('RB Allocation')
                displaycolumns.remove('RB Offset')    
            if "WLAN" not in group:
                displaycolumns.remove('Max Area (W/kg)')
            if values['-peak_sar-'] == "No":
                displaycolumns.remove('Peak SAR Location (x,y,z)')
            window['-data_table_1-'].ColumnsToDisplay = displaycolumns
            window['-data_table_2-'].ColumnsToDisplay = displaycolumns
            window['-data_table_1-'].Widget.configure(displaycolumns=displaycolumns)
            window['-data_table_2-'].Widget.configure(displaycolumns=displaycolumns)
        except:
            continue
                
        window["-data_table_1-"].update(values = data)
        xl.close()
        
    elif event == "Load Docx" and path_docx != "":
        window["-Error_Technology-"].update("")
        window["-FuckedUp-"].update("") 

        try:
            docx = Document(path_docx)  # Open docx.
        except:
            window["-Error_Technology-"].update("Please use a Docx file!")
            window["-FuckedUp-"].update("Dx")            
        docx_paragraphs = [para.text for para in docx.paragraphs if ((para.text).strip() not in ["", "\n"])] # Remove random whitespace and newlines from the paragraph list. # Load all paragraphs in the docx.
        docx_tables = docx.tables   # Load all tables in the docx.
        main_table = []   # Initialize the list that will hold data for the main window.
        probe_table = []  # Initialize the list that will hold data for the probe.
        dae_table = []    # Initialize the list that will hold data for the DAE.
        liquid_table = [] # Initialize the list that will hold the data for the liquid checker window.
        peak_sar_temp = [] # Initialize a temp list for holding the peak SAR location.
        '''
            Notes about this section:
            The plots contains tables that contains the data that we need to parse. There are four tables with data that we can parse from.
            The following data will be placed into a 2D array that will be combined with a later 2D array to form the list that the table from PySimpleGUI will use:
            1) Frequency
            2) Channel
            3) Group (NOTE: This is what SPEAG categorizes the current test as, technology wise)
            4) RF Exposure Condition
            5) 1-g Max Area Scan SAR
            6) 1-g Measured Zoom Scan SAR
            7) 10-g Measured Zoom Scan SAR
            
            Other pieces of data that it will grab:
            1) Probe SN
            2) Probe Cal Date
            3) DAE SN
            4) DAE Cal Date
            5) Relative Permittivity
            6) Conductivity
            7) Target Frequency for Liquid Check window
        '''
        start_of_tables = 0 # Determines the table that is having its data parsed.
        end_of_tables = 4
        sublist_start = 0 # Determines which index of the list holding the data (table_1) to extend it with more data.
        while start_of_tables < end_of_tables:
            for table_num in range(start_of_tables, len(docx_tables), 4):
                table = docx_tables[table_num]  # Define current table from docx.
                
                if start_of_tables == 0: # 'start_of_tables = 0' is the "Exposure Conditions" table on the plot.
                    split_freqch = (table.rows[1].cells[1].text).split()
                    permittivity = table.rows[0].cells[3].text                                    # Get relative permittivity.
                    conductivity = table.rows[1].cells[3].text                                    # Get conductivity.
                    frequency, channel = '{:.1f}'.format(float(split_freqch[0])), split_freqch[2] # Get frequency and channel.
                    group = (table.rows[2].cells[1].text).split()                                 # Get group.
                    test_distance = table.rows[3].cells[3].text                                   # Get test distance.
                    target = frequency                                                            # Get target frequency.
                    
                    main_table.append([test_distance,
                                       channel,
                                       frequency])
                    
                    liquid_table.append([target,
                                         permittivity,
                                         conductivity])
                    
                    del split_freqch    # Clear memory.
                    
                elif start_of_tables == 1: # 'start_of_tables = 1' is the "Hardware Setup" table on the plot.
                    split_probe_sncal = (table.rows[0].cells[1].text).split()
                    probe_sn = "".join(text for text in split_probe_sncal[0:3]) # Get probe serial number.
                    probe_cal = split_probe_sncal[len(split_probe_sncal)-1]     # Get probe calibration date.
                    probe_caldue = str(int(probe_cal[0:4]) + 1) + probe_cal[4:] # Get probe calibration due date. (1 year from calibration date)
                    
                    split_dae_sncal = (table.rows[1].cells[1].text).split()
                    dae_sn = " ".join(text for text in split_dae_sncal[0:2]) # Get DAE serial number.
                    dae_cal = split_dae_sncal[len(split_dae_sncal)-1]        # Get DAE calibration date.
                    dae_caldue = str(int(dae_cal[0:4]) + 1) + dae_cal[4:]    # Get DAE calibration due date. (1 year from calibration date)
                                
                    probe_table.append([probe_sn,
                                        probe_cal,
                                        probe_caldue])
                    dae_table.append([dae_sn,
                                      dae_cal,
                                      dae_caldue])
                    
                    del split_probe_sncal   # Clear memory.
                    del split_dae_sncal     # Clear memory.
                elif start_of_tables == 2: # 'start_of_tables = 2' is the "Scan Setup' table on the plot
                    if len(table.rows[0].cells) == 2: # Only area scan case.
                        peak_sar_temp.append("N/A")
                    if len(table.rows[0].cells) == 3: # Only one zoom scan case.
                        peak_sar = (table.rows[7].cells[2].text).split("|")[1]
                        for replace in ['[', ']']:
                            peak_sar = peak_sar.replace(replace,'')
                        peak_sar_temp.append(peak_sar)
                    if len(table.rows[0].cells) == 4: # Only two zoom scans case.
                        first_peak = (table.rows[7].cells[2].text).split("|")[1]
                        second_peak = (table.rows[7].cells[3].text).split("|")[1]
                        for bracket in ['[', ']']:
                            first_peak = first_peak.replace(bracket,'')
                            second_peak = second_peak.replace(bracket,'')
                        peak_sar_temp.append([first_peak, second_peak])
                elif start_of_tables == 3:                                 # 'start_of_tables = 3' is the "Measurement Results" table on the plot.
                    max_area_scan_1g = table.rows[1].cells[1].text         # Get max area scan's measured 1-g (W/kg).
                    max_area_scan_10g = table.rows[2].cells[1].text        # Get max area scan's measured 10-g (W/kg).
                    if len(table.rows[0].cells) == 2:                      # A table length of '2' means that there is only an area scan value.
                        zoom_meas_1g = "N/A"                               # There is no measured 1-g when only area scan.
                        zoom_meas_10g = "N/A"                              # There is no measured 10-g when only area scan.
                        power_drift = "N/A"                                # There is no power drift when only area scan.
                    elif len(table.rows[0].cells) == 3:                    # A table length of '3' means that there is one scan.
                        zoom_meas_1g = table.rows[1].cells[2].text         # Get zoom scans measured 1-g (W/kg).
                        zoom_meas_10g = table.rows[2].cells[2].text        # Get zoom scans measured 10-g (W/kg).
                        power_drift = table.rows[3].cells[2].text          # Get power drift (dB).
                    elif len(table.rows[0].cells) == 4:                    # A table length of '4' means that there are two scans.
                        first_zoom_meas_1g = table.rows[1].cells[2].text   # Get first zoom scan's measured 1-g (W/kg).
                        first_zoom_meas_10g = table.rows[2].cells[2].text  # Get first zoom scan's measured 10-g (W/kg).
                        first_power_drift = table.rows[3].cells[2].text    # Get first power drift (dB).
                        second_zoom_meas_1g = table.rows[1].cells[3].text  # Get second zoom scan's measured 1-g (W/kg).
                        second_zoom_meas_10g = table.rows[2].cells[3].text # Get second zoom scan's measured 10-g (W/kg).
                        second_power_drift = table.rows[3].cells[3].text   # Get second power drift (dB).
                    
                        zoom_meas_1g = first_zoom_meas_1g if first_zoom_meas_1g > second_zoom_meas_1g else second_zoom_meas_1g      # Determines which 1-g measured zoom scan to use.
                        zoom_meas_10g = first_zoom_meas_10g if first_zoom_meas_10g > second_zoom_meas_10g else second_zoom_meas_10g # Determines which 10-g measured zoom scan to use.
                        
                        # Determine which power drift to use from which zoom scan.
                        if first_power_drift == "N/A" or second_power_drift == "N/A":
                            power_drift = "N/A"
                        else:
                            if abs(float(first_power_drift)) > 0.20 and abs(float(second_power_drift)) > 0.20:
                                power_drift = "0.20"
                            elif first_zoom_meas_1g > second_zoom_meas_1g and first_zoom_meas_10g > second_zoom_meas_10g:
                                power_drift = "0.20" if abs(float(first_power_drift)) > 0.20 else first_power_drift
                            elif first_zoom_meas_1g < second_zoom_meas_1g and first_zoom_meas_10g < second_zoom_meas_10g:
                                power_drift = "0.20" if abs(float(first_power_drift)) > 0.20 else second_power_drift
                            else:
                                power_drift = "N/A"
                            
                    # Extend current sublist of 'table_1' with max area scan, 1-g measured, 10-g measured.
                    if sublist_start < len(main_table):
                        if "WLAN" in group:
                            main_table[sublist_start].extend(["{:.3f}".format(round(float(max_area_scan_1g), 3))])
                        
                        if (zoom_meas_1g != "N/A" and zoom_meas_10g != "N/A") and (power_drift != "N/A"):
                            main_table[sublist_start].extend(["{:.3f}".format(round(float(zoom_meas_1g), 3)),
                                                           "{:.3f}".format(round(float(zoom_meas_10g), 3)),
                                                           "{:.2f}".format(round(float(power_drift), 2))])
                        else:
                            main_table[sublist_start].extend([zoom_meas_1g, zoom_meas_10g, power_drift])
                    sublist_start += 1
            sublist_start = 0
            start_of_tables += 1

        # Add the peak sar location onto the main table list.
        sublist_start = 0
        for curr_peak in peak_sar_temp:         
            if type(curr_peak) is str:
                main_table[sublist_start].insert(len(main_table[sublist_start])-1, curr_peak)
            elif type(curr_peak) is list:
                if main_table[sublist_start][len(main_table[sublist_start])-3] > main_table[sublist_start][len(main_table[sublist_start])-2]:
                    main_table[sublist_start].insert(len(main_table[sublist_start])-1, curr_peak[0])
                elif main_table[sublist_start][len(main_table[sublist_start])-3] < main_table[sublist_start][len(main_table[sublist_start])-2]:
                    main_table[sublist_start].insert(len(main_table[sublist_start])-1, curr_peak[0])
                else:
                    main_table[sublist_start].insert(len(main_table[sublist_start])-1, curr_peak[0])
            else:
                main_table[sublist_start].insert(len(main_table[sublist_start])-1, "N/A")
            sublist_start += 1
        '''
            Notes about this section:
            This for-loop will iterate through each of the paragraphs of the plot in order to get the following information and place them in a 2D list for the PySimpleGUI table:
            1) Test Position (LHS Touch, LHS Tilt, RHS Touch, RHS Tilt, Back, Front, Edge Top, Edge Right, Edge Bottom, Edge Left)
            2) Band (GSM, WCDMA, LTE, FR1, WLAN, Bluetooth)
            3) RB Allocation (LTE and FR1 only)
            4) RB Offset (LTE and FR1 only)
        '''
        plot_data = []
        start_of_tables = 0
        index = 0
        index_equipment_liquid = 0
        plot_num = 0
        for paragraph_num in range(0, len(docx_paragraphs)):
            para_index = docx_paragraphs[paragraph_num] # Holds the current string in the paragraph.
            
            # Logic to add the plot number to the table and also add Probe, DAE, and Liquid Check data to their own respective tables.
            if paragraph_num % 6 == 0 or paragraph_num == 0:
                plot_num += 1
                plot_data.append([plot_num])
                probe_table[index_equipment_liquid].insert(0, plot_num)
                dae_table[index_equipment_liquid].insert(0, plot_num)
                liquid_table[index_equipment_liquid].insert(0, plot_num)
                
                sar_lab = (para_index[para_index.find("SAR Lab"):(para_index.find("SAR Lab")+len("SAR Lab")+3)]).strip()
                split_date_tested = (para_index[para_index.find("Date/Time:"):]).split() # In the plot it is listed as 'Date/Time: ####-##-##, ##:##'
                date_tested = split_date_tested[1] + " " + split_date_tested[2]          # Get date and time specifically.

                # Insert probe and dae data to its own table.
                if any(sar_lab not in equipment[index_equipment_liquid] for equipment in [probe_table, dae_table]):
                    probe_table[index_equipment_liquid].insert(1, sar_lab)
                    dae_table[index_equipment_liquid].insert(1, sar_lab)
                
                # Insert liquid data to its own table.                
                if date_tested not in liquid_table:
                    liquid_table[index_equipment_liquid].insert(1, sar_lab)
                    liquid_table[index_equipment_liquid].insert(2, date_tested)
    
                index_equipment_liquid += 1 # Move up one index for probe/DAE and liquid check tables.
            
            # Logic to get position.
            split_head = (docx_tables[start_of_tables].rows[2].cells[3].text).split()
            if any([pos in para_index for pos in ["CHEEK", "TILT", "BACK", "FRONT", "EDGE TOP", "EDGE RIGHT", "EDGE BOTTOM", "EDGE LEFT"]]):
                for pos in ["CHEEK", "TILT", "BACK", "FRONT", "EDGE TOP", "EDGE RIGHT", "EDGE BOTTOM", "EDGE LEFT"]:
                    position = para_index[para_index.find("{}".format(pos)):].strip()
                    if position in ["CHEEK", "TILT"]:
                        side = "Left" if "LeftHead" in split_head else "Right"
                        cheek_tilt = position[0] + position[1:].lower()
                        position = side + " " + cheek_tilt
                        break
                    elif position in ["BACK", "FRONT", "EDGE TOP", "EDGE RIGHT", "EDGE BOTTOM", "EDGE LEFT"]:
                        if position in ["BACK", "FRONT"]:
                            position = position[0] + position[1:].lower()
                        elif position in ["EDGE TOP", "EDGE RIGHT", "EDGE BOTTOM", "EDGE LEFT"]:
                            edge_split = position.split()    # Split 'EDGE' and side into list.
                            edge = edge_split[0][0] + edge_split[0][1:].lower() # Logic to get 'EDGE' and lowercase everything after the first character.
                            side = edge_split[1][0] + edge_split[1][1:].lower() # Logic to get the side and lowercase everything after the first character.
                            position = ' '.join([edge, side])
                        break
                
                # Logic to get mode that is used.
                '''
                    Notes about this section:
                    Generally, the following are the modes used for each technology:
                        NOTE: There are other modes that we can technically test. These modes, fortunately, are almost never tested. 
                              If there is ever a time were it is common to test the other modes, I'll add them in.
                    
                    GSM Mode   = GPRS [#] Slot(s)
                                 Where # is the slot number, which ranges from 1 to 4.
                    WCDMA Mode = Rel 99
                    LTE Mode   = QPSK
                    FR1 Mode   = DFT-s-OFDM π/2 BPSK
                    WLAN Modes = 802.11[mode]. 
                                 Where mode is dependent on frequency range: 
                                 2.4 GHz = b, g, or n
                                 5 GHz   = a, n, ac, or ax
                                 6 GHz   = n, ac, or ax
                    Bluetooth  = GFSK (BDR)
                '''
                if "GSM" in group:
                    find_slots = para_index[para_index.find("TN"):para_index.find(")")].strip()
                    slots_on_plot = find_slots[3:]
                    if slots_on_plot == "0":
                        slot_num = "1"
                    elif slots_on_plot == "0-1":
                        slot_num = "2"
                    elif slots_on_plot == "0-1-2":
                        slot_num = "3"
                    elif slots_on_plot == "0-1-2-3":
                        slot_num = "4"
                    mode = "GPRS {} {}".format(slot_num, "Slot" if int(slot_num) == 1 else "Slots")
                elif "WCDMA" in group:
                    mode = "Rel. 99"
                elif "LTE-FDD" in group or "LTE-TDD" in group:
                    mode = "QPSK"
                elif "FR1" in group:
                    mode = "DFT-s-OFDM π/2 BPSK"
                elif "WLAN" in group:
                    for check_letter in ["802.11b", "802.11g", "802.11n", "802.11a", "802.11ac", "802.11ax", "802.11be"]:
                        if check_letter in para_index:
                            mode = para_index[para_index.find("802.11"):(para_index.find("802.11")+len("{}".format(check_letter)))].strip()
                elif "Bluetooth" in group:
                    mode = "GFSK (BDR)"

                # Logic to get the RF exposure condition based on test distance.
                '''
                    Notes about this section:
                    The following test distances are what's typically used:
                        NOTE: All test distances are in millimeters (mm).

                    0.00  = Head or Extremity (Generally)
                    5.00  = Body-Worn/Hotspot (NOTE: This test distance is for a certain client)
                    10.00 = Hotspot
                    15.00 = Body-Worn
                    
                    Anything other than these test distances are usually counted as "Extremity".
                '''
                if "0.00" == main_table[index][0]:
                    rf_exposure_condition = "Head" if any(pos in position for pos in ["Left Cheek", "Left Tilt", "Right Cheek", "Right Tilt"]) else "Extremity"
                elif "5.00" == main_table[index][0]:
                    rf_exposure_condition = "Body-Worn/Hotspot"
                elif "10.00" == main_table[index][0]:
                    rf_exposure_condition = "Hotspot"
                elif "15.00" == main_table[index][0]:
                    rf_exposure_condition = "Body-Worn"
                del main_table[index][0] # Remove test distance flag.
                
                plot_data[index].insert(1, rf_exposure_condition) # Insert RF exposure condition into plot data list.
                plot_data[index].insert(2, mode)                  # Insert Mode into plot data list.
                plot_data[index].insert(3, position)              # Insert Test position into plot data list.
                
                # Go to the next paragraph.
                if start_of_tables < len(docx_tables)-4:
                    start_of_tables += 4
            
            # Logic to get Bluetooth, GSM/PCS and WCDMA.
            if any(technology in para_index for technology in ["GSM", "PCS", "WCDMA", "Bluetooth"]):
                # Get the techonology that is used.
                gsm_pcs_tech = "GSM" if para_index.find("GSM") != -1 else "PCS"
                if "GSM" in para_index:
                    find_colon_gsm_850_900 = para_index[para_index.find("GSM")+7]
                    if find_colon_gsm_850_900 == ":":
                        # Band is 850 or 900.
                        technology = para_index[para_index.find("GSM"):para_index.find("GSM")+7]
                    else:
                        # Band is 1800 or 1900
                        technology = para_index[para_index.find("GSM"):para_index.find("GSM")+8]
                elif "PCS" in para_index:
                    find_colon_gsm_850_900 = para_index[para_index.find("PCS")+7]
                    if find_colon_gsm_850_900 == ":":
                        # Band is 850 or 900.
                        technology = para_index[para_index.find("PCS"):para_index.find("PCS")+7]
                    else:
                        # Band is 1800 or 1900.
                        technology = para_index[para_index.find("PCS"):para_index.find("PCS")+8]
                elif "WCDMA" in para_index:
                    # Get techonology and band 
                    wcdma, band = para_index[para_index.find("WCDMA"):para_index.find("WCDMA")+5], para_index[para_index.find("Band"):para_index.find("Band")+6]
                    technology =  wcdma + " " + band
                elif "Bluetooth" in para_index:
                    bluetooth = para_index[para_index.find("Bluetooth"):para_index.find("Bluetooth")+len("Bluetooth")]
                    technology = bluetooth
                    
                # GSM / WCDMA / Wi-Fi / Bluetooth don't have RBs
                main_table[index].insert(2, "N/A") # Fill table with 'N/A' for RB Allocation.
                main_table[index].insert(3, "N/A") # Fill table with 'N/A' for RB Offset.
                main_table[index].insert(4, "N/A") # Fill table with 'N/A' for Max Area Scan.
            # Logic for Wi-Fi.
            elif any(wifi in para_index for wifi in ["Wi-Fi", "Wi-fi", "WI-FI", "Wifi", "WLAN", "UNII", "U-NII"]):
                main_table[index].insert(2, "N/A") # Fill table with 'N/A' for RB Allocation.
                main_table[index].insert(3, "N/A") # Fill table with 'N/A' for RB Offset.
            # Logic to get LTE and 5G NR (FR1).
            elif "LTE" in para_index or "5G NR" in para_index:
                # Get the technology that is used.
                tech = para_index[para_index.find("LTE"):para_index.find("LTE")+3] if "LTE" in para_index else para_index[para_index.find("5G NR")+2:para_index.find("5G NR")+5]
                
                # Get the band that is used.
                band = para_index[para_index.find("Band"):para_index.find("Band")+7] if para_index[para_index.find("Band")+7].isdigit() else para_index[para_index.find("Band"):para_index.find("Band")+6]
                
                # Merge technology and band into one string.
                technology = tech + " " + band

                # Get the RB position.
                if para_index.find("Low") != -1:
                    rb_position = para_index[para_index.find("Low"):para_index.find("Low")+3]
                elif para_index.find("Mid") != -1:
                    rb_position = para_index[para_index.find("Mid"):para_index.find("Mid")+3]
                elif para_index.find("High") != -1:
                    rb_position = para_index[para_index.find("High"):para_index.find("High")+4]
                
                # Is the current plot 1 RB, 50% RB, or 100% RB?
                if para_index.find("1 RB") != -1:
                    num_rb = para_index[para_index.find("1 RB"):para_index.find("1 RB")+4]
                elif para_index.find("50% RB") != -1:
                    num_rb = para_index[para_index.find("50% RB"):para_index.find("50% RB")+6]
                elif para_index.find("100% RB") != -1:
                    num_rb = para_index[para_index.find("100% RB"):para_index.find("100% RB")+7]
                
                # This section is the logic to get the NRB (Number of Resource Blocks)
                check_half_rb = para_index.find("50%") != -1     # Check if the current plot is for 50% RB
                check_full_rb = para_index.find("100%") != -1    # Check if the current plot is for 100% RB
                find_bw = str(para_index[para_index.find("RB,")+3:para_index.find("MHz")-1]).strip() # Get the bandwidth number.
                bw_hz = float(find_bw) * pow(10, 6) # Bandwidth in Hz
                # This section is the logic to get the RB allocation and RB offset for LTE.
                if "{} MHz".format(str(find_bw)) in para_index and ("LTE" in para_index):
                    # When the current bandwidth for the band is less than 3 MHz.
                    if bw_hz / pow(10,6) < 3:
                        size_guard_band = bw_hz * 0.001                                              # Guardband = 10% of BW (In Hz)
                        single_slot = 12 * 15                                                        # 12 subcarriers * 15 kHz subcarrier spacing = 180 kHz size of 1 slot.
                        nRB_ref = math.floor((size_guard_band)/(single_slot))                        # Reference NRB which is the BW / slot size. I.e. for 1.4 MHz 1400 kHz / 180 kHz = floor(7.77) = 7.
                        guard_band = (size_guard_band - ((nRB_ref * single_slot) - single_slot)) / 2 # Calculate guardband in kHz.
                        used_bw = size_guard_band - (guard_band * 2)                                 # Calculate the usable bandwidth, (BW w/guard - (guard band size * 2)) = usable BW (kHz)
                        full_rb = math.floor(used_bw / single_slot)                                  # Calculate the full number of RBs that users can allocate.
                    # When the current bandwidth is in [3, 5, 10, 15, 20] (MHz)
                    else:
                        size_guard_band = bw_hz * 0.1               # Guardband = 10% of BW (in Hz).
                        single_slot = 12 * 15 * pow(10, 3)          # Size of single slot (in Hz).
                        used_bw = bw_hz - size_guard_band           # The used bandwidth is the BW (in Hz) - guard band (in Hz).
                        full_rb = math.floor(used_bw / single_slot) # To get the number of RBs: Used bandwidth / single slot.

                    mhz = pow(10,6) # 10^6 for MHz convertion.
                    
                    # Get the RB allocation and offset. This is dependent on the RB position and what percentage of the RB is being allocated.
                    # When 50% of the RBs are allocated.
                    if check_half_rb:
                        if any(bw_hz / pow(10,6) == bw for bw in [3, 15]):
                            rb_allocation = "8" if bw_hz / pow(10,6) == 3 else "36"
                        else:
                            rb_allocation = str(math.floor(full_rb / 2)) # Half of the number of RBs is being allocated.
                            
                        rb_offset = rb_positions[str(int(bw_hz))]["LTE"][num_rb][rb_position]            # Grab RB offset from dictionary. 
                    # When 100% of the RBs are allocated.
                    elif check_full_rb:
                        rb_allocation = str(full_rb) # 100% of the RB is being allocated.
                        rb_offset = "0"              # There is no offset when 100% RB is being allocated, otherwise band-edge.
                    # When the amount of RBs allocated is 1.
                    else:
                        rb_allocation = "1" # 1 RB is being allocated.
                        rb_offset = rb_positions[str(int(bw_hz))]["LTE"][num_rb][rb_position]            # Grab RB offset from dictionary.
                    main_table[index].insert(2, rb_allocation) # Insert number of allocated RBs.
                    main_table[index].insert(3, rb_offset)     # Insert offset for RBs.
                    main_table[index].insert(4, "N/A")         # Fill table with 'N/A' for Max Area Scan.
                # This is the logic to get 
                elif "{} MHz".format(str(find_bw)) in para_index and ("5G NR" in para_index):
                    find_scs = str(para_index[para_index.find("kHz")-3:para_index.find("kHz")-1]).strip() # Find subcarrier spacing in the plot.
                    scs_hz = float(find_scs) * pow(10,3)                                                  # SCS in Hz.
                    full_rb = int(nr_nrbs[str(int(scs_hz))][str(int(bw_hz))])                             # Get number of RBs.
                    
                    # Get the RB allocation and offset. This is dependent on the RB position and what percentage of the RB is being allocated.
                    # When 50% of the RBs are allocated.
                    if check_half_rb:
                        rb_allocation = str(math.floor(full_rb / 2)) # Half of the number of RBs is being allocated.
                        rb_offset = rb_positions[str(int(bw_hz))]["NR"][str(int(scs_hz))][num_rb][rb_position] # Grab RB offset from dictionary. 
                    # When 100% of the RBs are allocated.
                    elif check_full_rb:
                        rb_allocation = str(full_rb) # 100% of the RB is being allocated.
                        rb_offset = "0"              # There is no offset when 100% RB is being allocated, otherwise band-edge.
                    # When the amount of RBs allocated is 1.
                    else:
                        rb_allocation = "1" # 1 RB is being allocated.
                        rb_offset = rb_positions[str(int(bw_hz))]["NR"][str(int(scs_hz))][num_rb][rb_position] # Grab RB offset from dictionary.
                        #else:
                        #    rb_offset = str(full_rb - 1) 
                    main_table[index].insert(2, rb_allocation) # Insert number of allocated RBs.
                    main_table[index].insert(3, rb_offset)     # Insert offset for RBs.
                    main_table[index].insert(4, "N/A")         # Fill table with 'N/A' for Max Area Scan.
            
            # If the current plot has ONLY AREA SCANS, increment. (No zoom scan data).
            if len(plot_data[index]) == 4 and (paragraph_num % 6 == 0 or paragraph_num == 0):
                index += 1

        df_plot = (pd.DataFrame(main_table)).values.tolist()

        # Merge the 'Plot #' and 'Position' with the 'Channel #', 'Frequency', 'RB allocation' (if applicable), 'RB offset' (if applicable), '1-g Measured SAR', and '10-g Measured SAR'.
        for index in range(0, len(df_plot)):
            df_plot[index] = plot_data[index] + df_plot[index]
            
        data_plot = df_plot # Used for comparison purposes.
        
        '''
            Notes about this section:
            The purpose of this section is to hide the columns on the table that are unnessesary when checking certain technologies. I.E. RB Allocation, RB Offset, and Max Area Scan.
        '''
        try:
            displaycolumns = deepcopy(COL_HEADINGS) # Creating a deepcopy as to not override the original.
            if any(technology in group for technology in ["GSM", "WCDMA", "WLAN", "Bluetooth"]):
                displaycolumns.remove('RB Allocation')
                displaycolumns.remove('RB Offset')
            if "WLAN" not in group:
                displaycolumns.remove('Max Area (W/kg)')
            if values["-peak_sar-"] == "No":
                displaycolumns.remove('Peak SAR Location (x,y,z)')
            window['-data_table_1-'].ColumnsToDisplay = displaycolumns
            window['-data_table_2-'].ColumnsToDisplay = displaycolumns
            window['-data_table_1-'].Widget.configure(displaycolumns=displaycolumns)
            window['-data_table_2-'].Widget.configure(displaycolumns=displaycolumns)
        except:
            continue

        window["-data_table_2-"].update(values = df_plot) # Update bottom table.
    # NOTE: !!!!!!!!!!! OPTIONAL: If I have the time and willpower, try to figure this section out !!!!!!!!!!!
    # elif event == "-data_table_1-":
    #     current_click = values["-data_table_1-"][0]
    #     data_selected = data_excel[current_click]
    #     print(data_selected)
    # elif event == "-data_table_2-":
    #     print(values["-data_table_2-"])
    # NOTE: !!!!!!!!!!! OPTIONAL: If I have the time and willpower, try to figure this section out !!!!!!!!!!!
    
    elif event == '-hide_unhide-' and values["-hide_unhide_list-"] != "":
        value = values["-hide_unhide_list-"]
        flag = 0
        
        displaycolumns = deepcopy(COL_HEADINGS) # Creating a deepcopy as to not override the original.
        if value == "Only 1g/10g values":
            for i in (list(set(COL_HEADINGS)-set(['Plot #', 'Max Area (W/kg)', '1-g Meas. (W/kg)', '10-g Meas. (W/kg)', 'Peak SAR Location (x,y,z)', 'Power Drift (dB)']))):
                displaycolumns.remove(i)
        elif value == "No 1g/10g values":
            for i in (list(set(COL_HEADINGS)-set(['Plot #', 'RF Exposure Condition', 'Mode', 'Test Position', 'Ch #.', 'Freq. (MHz)', 'RB Allocation', 'RB Offset']))):
                displaycolumns.remove(i)
        else:
            displaycolumns = deepcopy(COL_HEADINGS)
        window['-data_table_1-'].ColumnsToDisplay = displaycolumns
        window['-data_table_2-'].ColumnsToDisplay = displaycolumns
        window['-data_table_1-'].Widget.configure(displaycolumns=displaycolumns)
        window['-data_table_2-'].Widget.configure(displaycolumns=displaycolumns)
        
        
    # Open the liquid check window if it has not been opened already.
    elif event == '-liquid_check-' and not window_liquid:
        window_liquid = make_win2()
    # Logic to calculate the relative permitivity and conductivity given a target frequency.
    elif event == 'Calculate' and values["-target_1-"].strip() != '' and values["-file_1-"].strip() != '':
        target = float(values["-target_1-"])
        
        try:
            file_1 = values["-file_1-"]
            myfile = open(file_1, "rt")
            lines = myfile.readlines()
            
            permcondCalcSolo(target, lines)
        
            myfile.close()
        except:
            continue

        myfile.close()
    # Open the Comparator window if it has not been open already.
    elif event == 'Load Parameters' and path_docx != "":
        for row, index in zip(liquid_table, range(0, len(liquid_table))):
            target = round(float(row[3]), 1) # Get target frequency.
            
            # Logic to format the date to search for .prn file.
            # Date on PRN ((Year: ####)-(Month: abc)-(Day: ##))
            date_tested = row[2]
            months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
            date_prn = (date_tested[:date_tested.find(",")]).split("-")
            date_prn[1] = months[int(date_prn[1])-1]
            date_prn = "-".join(date_prn)
            
            sar_lab = liquid_table[index][1] # Need SAR Lab to search for PRN file.
            
            prn_file = find_prn(target, sar_lab, date_prn, months) # Get PRN file.
            
            try:
                myfile = open(prn_file, "rt")
                lines = myfile.readlines()
                
                rpermittivity, conductivity = permcondCalcTable(target, lines)

                # NOTE: Permittivity is formatted - #.##, Conductivity is formatted - #.## or 0.###.
                liquid_table[index].append("{}".format(round(float(rpermittivity), 1)))     # Add and format relative permittivity to table.
                if conductivity[0] == "0":
                    if len("{}".format(round(float(conductivity), 3))) < 5:
                        liquid_table[index].append("{}0".format(round(float(conductivity), 3)))  # Add and format conductivity to table if the conductivity < 1.
                    else:
                        liquid_table[index].append("{}".format(round(float(conductivity), 3)))  # Add and format conductivity to table if the conductivity < 1.
                elif len("{}".format(round(float(conductivity), 2))) == 3:                  
                    liquid_table[index].append("{}0".format(round(float(conductivity), 2))) # Add and format conductivity to table if the rounding only produces 2 digits.
                else:
                    liquid_table[index].append("{}".format(round(float(conductivity), 2)))  # Add and format conductivity to table otherwise.
                myfile.close()
                
                # Compare plot and PRN data.
                rperm_plot = row[4] # Assign relative permittivity from plot.
                cond_plot =  row[5] # Assign conductivity from plot.
                rperm_prn =  row[6] # Assign relative permittivity calculated from PRN.
                cond_prn =   row[7] # Assign conductivity calculated from PRN.
                if (rperm_plot == rperm_prn) and (cond_plot == cond_prn):
                    liquid_table[index].append("Y")
                else:
                    liquid_table[index].append("N")
            except:
                continue
        window['-liquid_table_1-'].update(values=liquid_table)
    elif event == "-compare-" and not window_compare:   
        window_compare = make_win3()
    elif event == "Prepare to Be Sad":
        window3 = window['-data_table_3-']

        # Compare data from excel and plots.
        match_xlsx_docx, compare_data = append_data([], data_excel, data_plot)
                
        # # Logic to show on the table if there is a mismatch in the data comparisons.
        # for row in match_xlsx_docx:
        #     if "N" in row[4:]:
        #         window3.update(row_colors=[(match_xlsx_docx.index(row), "red")])
        #     else:
        #         window3.update(row_colors=[(match_xlsx_docx.index(row), "")])
                
        window3.update(values = compare_data) # Insert data into table.
    # Open Equipment window if it has not been open yet.
    elif event == "-equipment-" and not window_equipment:
        window_equipment = make_win4()
    # Logic to load the equipment data into the Equipment window.
    elif event == "-load_equipment-":
        try:
            window["-equipment_table_1-"].update(values=probe_table) # Insert probe data to top table.
            window["-equipment_table_2-"].update(values=dae_table)   # Insert dae data to bottom table.
        except:
            window["-equipment_table_1-"].update(values="")
            window["-equipment_table_2-"].update(values="")
    # Errors messages.
    elif (event == "Load" or (event == "Load Excel" and path == "")) or (event == "Load Excel" and tech == "") or (event == "Load Docx" and path_docx == ""):
        # Display text to user if a step was performed before a certain other step.
        if event == "Load" or (event == "Load Excel" and path == ""):
            error = "Please load an excel file."
        elif event == "Load Excel" and (tech == "" or values["-tech_1-"] == ""):
            error = "Please load a technology."
        elif event == "Load Docx" and path_docx == "":
            error = "Please load a docx file."
        else:
            error = ""
        window["-FuckedUp-"].update(error)
window.close()
